from __future__ import annotations

import copy
import re
from io import BytesIO
from typing import Dict, List, Tuple, TypedDict, Union, Optional

from service.ragfusion.core.document.document import Document


class MarkdownHeaderTextSplitter:
    """Splitting markdown files based on specified headers."""

    def __init__(
            self,
            headers_to_split_on: List[Tuple[str, str]],
            return_each_line: bool = False,
            strip_headers: bool = True,
            combine_content_with_header: bool = False
    ):
        """Create a new MarkdownHeaderTextSplitter.

        Args:
            headers_to_split_on: Headers we want to track
            return_each_line: Return each line w/ associated headers
            strip_headers: Strip split headers from the content of the chunk
            combine_content_with_header: Default by read line and split, it can combine content with same headers
        """
        # Output line-by-line or aggregated into chunks w/ common headers
        self.return_each_line = return_each_line
        # Given the headers we want to split on,
        # (e.g., "#, ##, etc") order by length
        self.headers_to_split_on = sorted(
            headers_to_split_on, key=lambda split: len(split[0]), reverse=True
        )
        # Strip headers split headers from the content of the chunk
        self.strip_headers = strip_headers
        # combine content with same headers
        self.combine_content_with_header = combine_content_with_header

    def aggregate_lines_to_chunks(self, lines: List[LineType]) -> List[Document]:
        """Combine lines with common metadata into chunks
        Args:
            lines: Line of text / associated header metadata
        """
        aggregated_chunks: List[LineType] = []

        for line in lines:
            if (
                    aggregated_chunks
                    and aggregated_chunks[-1]["metadata"] == line["metadata"]
            ):
                # If the last line in the aggregated list
                # has the same metadata as the current line,
                # append the current content to the last lines's content
                aggregated_chunks[-1]["content"] += "  \n" + line["content"]
            elif (
                    aggregated_chunks
                    and aggregated_chunks[-1]["metadata"] != line["metadata"]
                    # may be issues if other metadata is present
                    and len(aggregated_chunks[-1]["metadata"]) < len(line["metadata"])
                    and aggregated_chunks[-1]["content"].split("\n")[-1][0] == "#"
                    and not self.strip_headers
            ):
                # If the last line in the aggregated list
                # has different metadata as the current line,
                # and has shallower header level than the current line,
                # and the last line is a header,
                # and we are not stripping headers,
                # append the current content to the last line's content
                aggregated_chunks[-1]["content"] += "  \n" + line["content"]
                # and update the last line's metadata
                aggregated_chunks[-1]["metadata"] = line["metadata"]
            else:
                # Otherwise, append the current line to the aggregated list
                aggregated_chunks.append(line)

        return [
            Document(page_content=chunk["content"], metadata=chunk["metadata"])
            for chunk in aggregated_chunks
        ]

    def split_text(self, text: str) -> List[Document]:
        """Split markdown file
        Args:
            text: Markdown file"""

        # Split the input text by newline character ("\n").
        lines = text.split("\n")
        # Final output
        lines_with_metadata: List[LineType] = []
        # Content and metadata of the chunk currently being processed
        current_content: List[str] = []
        current_metadata: Dict[str, str] = {}
        # Keep track of the nested header structure
        # header_stack: List[Dict[str, Union[int, str]]] = []
        header_stack: List[HeaderType] = []
        initial_metadata: Dict[str, str] = {}

        in_code_block = False
        opening_fence = ""

        for line in lines:
            stripped_line = line.strip()

            if not in_code_block:
                # Exclude inline code spans
                if stripped_line.startswith("```") and stripped_line.count("```") == 1:
                    in_code_block = True
                    opening_fence = "```"
                elif stripped_line.startswith("~~~"):
                    in_code_block = True
                    opening_fence = "~~~"
            else:
                if stripped_line.startswith(opening_fence):
                    in_code_block = False
                    opening_fence = ""

            if in_code_block:
                current_content.append(stripped_line)
                continue

            # Check each line against each of the header types (e.g., #, ##)
            for sep, name in self.headers_to_split_on:
                # Check if line starts with a header that we intend to split on
                if stripped_line.startswith(sep) and (
                        # Header with no text OR header is followed by space
                        # Both are valid conditions that sep is being used a header
                        len(stripped_line) == len(sep) or stripped_line[len(sep)] == " "
                ):
                    # Ensure we are tracking the header as metadata
                    if name is not None:
                        # Get the current header level
                        current_header_level = sep.count("#")

                        # Pop out headers of lower or same level from the stack
                        while (
                                header_stack
                                and header_stack[-1]["level"] >= current_header_level
                        ):
                            # We have encountered a new header
                            # at the same or higher level
                            popped_header = header_stack.pop()
                            # Clear the metadata for the
                            # popped header in initial_metadata
                            if popped_header["name"] in initial_metadata:
                                initial_metadata.pop(popped_header["name"])

                        # Push the current header to the stack
                        header: HeaderType = {
                            "level": current_header_level,
                            "name": name,
                            "data": stripped_line[len(sep):].strip(),
                        }
                        header_stack.append(header)
                        # Update initial_metadata with the current header
                        initial_metadata[name] = header["data"]

                    # Add the previous line to the lines_with_metadata
                    # only if current_content is not empty
                    if current_content:
                        lines_with_metadata.append(
                            {
                                "content": "\n".join(current_content),
                                "metadata": current_metadata.copy(),
                            }
                        )
                        current_content.clear()

                    if not self.strip_headers:
                        current_content.append(stripped_line)

                    break
            else:
                if stripped_line:
                    current_content.append(stripped_line)
                elif current_content:
                    lines_with_metadata.append(
                        {
                            "content": "\n".join(current_content),
                            "metadata": current_metadata.copy(),
                        }
                    )
                    current_content.clear()

            current_metadata = initial_metadata.copy()

        if current_content:
            lines_with_metadata.append(
                {"content": "\n".join(current_content), "metadata": current_metadata}
            )

        # lines_with_metadata has each line with associated header metadata
        # aggregate these into chunks based on common metadata
        if not self.return_each_line:
            chunks = self.aggregate_lines_to_chunks(lines_with_metadata)
        else:
            chunks = [
                Document(page_content=chunk["content"], metadata=chunk["metadata"])
                for chunk in lines_with_metadata
            ]

        if self.combine_content_with_header:
            order_chunks = {
                tuple(doc.metadata.items()): [d.page_content for d in chunks if tuple(d.metadata.items()) == tuple(doc.metadata.items())]
                for doc in chunks
            }
            return [
                Document(
                    page_content="\n".join(chunk_texts),
                    metadata={key: value.rstrip("：") for key, value in header_tuple}
                )
                for header_tuple, chunk_texts in order_chunks.items()
            ]
        return chunks



class BaseHeaderSplitter:
    """Splitting headers"""

    def __init__(
            self,
            return_each_line: bool = False,
            strip_headers: bool = True,
            split_headers_re: Optional[str] = r'^\d+(?:\.\d+)',  # r'^\d+(?:\.\d+)'
            add_title_in_content: bool = False,
    ):
        # Output line-by-line or aggregated into chunks w/ common headers
        self.return_each_line = return_each_line
        # Strip headers split headers from the content of the chunk
        self.strip_headers = strip_headers
        # Split headers by re, default like '1.2.3 Summary'
        self.split_headers_re = split_headers_re
        # Save headers info in content for retrieve
        self.add_title_in_content = add_title_in_content

    def _word_headers_extractor(
            self,
            document_parts,
    ) -> List[Document]:
        """Splitting headers by specified word documents type"""

        # import docx
        #
        # try:
        #     document_parts = docx.Document(file_or_buffer)
        # except Exception as e:
        #     raise RuntimeError(f"Error loading {file_or_buffer}") from e

        # Record the newest title sequence
        heading_counters = {
            "Heading 1": None, "Heading 2": None,
            "Heading 3": None, "Heading 4": None
        }
        # Saving the sequence of documentation
        documents: List[Document] = []
        # Saving the temporary content
        content: str = ""

        for paragraph in document_parts.paragraphs:
            paragraph_style = paragraph.style.name
            if paragraph_style in heading_counters:
                copy_headers = copy.deepcopy(heading_counters)
                if content and self.add_title_in_content:
                    documents.append(Document(page_content=f"当前的段落层级结构是：" + '<{}>'.format('>-<'.join(str(v) for v in copy_headers.values() if v is not None)) + "\n" + content,
                                              metadata={k: v for k, v in copy_headers.items() if
                                                        v is not None}))
                else:
                    documents.append(Document(page_content=content,
                                              metadata={k: v for k, v in copy_headers.items() if
                                                        v is not None}))
                content = ""

                # Reset counters for lower level headings
                heading_counters[paragraph_style] = paragraph.text
                if paragraph_style == "Heading 1":
                    heading_counters["Heading 2"] = None
                    heading_counters["Heading 3"] = None
                    heading_counters["Heading 4"] = None
                elif paragraph_style == "Heading 2":
                    heading_counters["Heading 3"] = None
                    heading_counters["Heading 4"] = None
                elif paragraph_style == "Heading 3":
                    heading_counters["Heading 4"] = None

                # Append the current heading paragraph
                if not self.strip_headers:
                    documents.append(Document(page_content=paragraph.text,
                                              metadata={k: v for k, v in copy.deepcopy(heading_counters).items() if
                                                        v is not None}))

            # Accumulate non-heading content
            else:
                content += f"\n{paragraph.text}"

        # Append any remaining content as the last document
        if content:
            documents.append(Document(page_content=content,
                                      metadata={k: v for k, v in copy.deepcopy(heading_counters).items() if
                                                v is not None}))

        return documents

    def _text_headers_extractor(self, texts: List[Union[str, Document]] = None) -> List[Document]:
        """Splitting headers by re"""

        text = '\n'.join((
            str(item)
            if isinstance(item, str)
            else item.page_content
            for item in texts
        ))

        # Saving the sequence of documentation
        documents: List[Document] = []
        # Saving the temporary content
        content: str = ""
        # For comparing current new title and save
        old_title: str = ""
        for line in text.split("\n"):
            if not line:
                continue

            if re.match(self.split_headers_re, line.strip()):
                current_title = line.strip()

                # organize the document
                page_content = f"{old_title}\n" + content if content and self.add_title_in_content else content
                if page_content:
                    documents.append(Document(page_content=page_content, metadata={"Header": old_title}))
                content = ""

                # add title additional chunk
                if not self.strip_headers:
                    documents.append(Document(page_content=current_title, metadata={"Header": current_title}))
                old_title = current_title
            else:
                content += f"\n{line}"

        if content:
            documents.append(Document(page_content=content, metadata={"Header": old_title}))

        return documents


class WordHeaderTextSplitter(BaseHeaderSplitter):
    """Splitting Word files based on specified headers.

       Only file with title format can be extract.
    """

    def __init__(
            self,
            **kwargs
    ):
        """Create a new WordHeaderTextSplitter.

        Only support the type of docx file
        """
        super().__init__(**kwargs)

    def split_text(self, file_path: Union[str, BytesIO]) -> List[Document]:
        """Split Word file path
        Args:
            file_path: Word file path"""
        # if splite_headers_re is exist, it means use re match to split header,
        # otherwise use word title match
        import docx

        try:
            document_parts = docx.Document(file_path)
        except Exception as e:
            raise RuntimeError(f"Error loading {file_path}") from e

        chunks = self._word_headers_extractor(document_parts)
        return chunks


class TextHeaderSplitter(BaseHeaderSplitter):
    """Splitting text or document object based on re.

       Default extraction of types with "digit-point" format.
    """

    def __init__(
            self,
            split_headers_type: Optional[int] = 0,
            **kwargs
    ):
        """Create a new TextHeaderSplitter.

            :param split_headers_type: self-defined split with choose type

        """
        split_headers_re_map = {
            0: r"^\d+(?:\.\d+)",  # x.x.x.x
            1: r"^\(\d+\).*",  # (xx)
        }
        self.split_headers_type = split_headers_re_map.get(split_headers_type)
        super().__init__(split_headers_re=self.split_headers_type, **kwargs)

    def split_text(self, texts: List[Union[str, Document]]) -> List[Document]:
        """Split text
        Args:
            texts: list of text or document"""
        if not all(isinstance(item, (str, Document)) for item in texts):
            raise ValueError("Only can support type of Str or Document object.")

        return self._text_headers_extractor(texts)


class LineType(TypedDict):
    """Line type as typed dict."""

    metadata: Dict[str, str]
    content: str


class HeaderType(TypedDict):
    """Header type as typed dict."""

    level: int
    name: str
    data: str
